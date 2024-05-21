import streamlit as st
import replicate
from dotenv import load_dotenv
import os
import json
from streamlit_extras.let_it_rain import rain
import base64
import whois
from io import BytesIO
from docx import Document

# Load environment variables from .env file
load_dotenv()

# Get the Replicate API token from environment variables
REPLICATE_API_TOKEN = os.getenv("REPLICATE_API_TOKEN")

if not REPLICATE_API_TOKEN:
    st.error("REPLICATE_API_TOKEN is not set in the environment variables.")
else:
    # Set the Replicate API token in the environment
    os.environ["REPLICATE_API_TOKEN"] = REPLICATE_API_TOKEN

    # Function to check if domain exists
    def check_domain_exists(domain_name):
        try:
            domain_info = whois.whois(domain_name)
            return domain_info.domain_name is not None
        except Exception as e:
            return False

    # Sidebar for input fields and buttons
    with st.sidebar:
        st.sidebar.title("Eureka 💡")
        st.sidebar.caption("Your AI-Powered App Idea Generator")
        st.sidebar.markdown("\n\n\n")
        user_prompt = st.text_input("Enter a keyword, idea, or area of interest to start 🏁", key="user_prompt")

        # Add buttons
        generate_button = st.button("Generate My App Idea 🚀")
        lucky_button = st.button("Surprise Me! 🎲")

        # Check which button was pressed
        if generate_button:
            prompt_to_use = user_prompt
        elif lucky_button:
            prompt_to_use = "Generate a random app idea that would be useful"
        else:
            prompt_to_use = None

    if prompt_to_use is None or prompt_to_use.strip() == "":
        st.markdown("Get started by entering a keyword, idea, or area of interest and click 'Generate My App Idea 🚀' to receive an AI-generated app concept. Or click 'Surprise Me! 🎲' to get a random app idea 💡")

    if prompt_to_use:

        # Prepare the prompt template for JSON output
        prompt_template = f"""system
            You are a creative and insightful brand consultant and app idea generator. I will provide you with an idea or an area of focus, and you will generate an AI app concept around it. 
            You will generate a clever name for the app with puns, an emoji to match the description of the app, a small tagline for the app, 
            a description of what the app will do, and then appropriate descriptions as I mention below. Also give a background color appropriate for the app and make it in dark theme.
            The output should be in JSON format:
                "Name": "<App Name>.ai",
                "Tagline": "<Tagline>",
                "Description": "<Description>",
                "Background_Color": "<Background_Color_dark_mode>",
                "Emoji": "<Emoji>",
                "Problem": "<Problem>",
                "Solution": "<Solution>",
                "Features": ["Feature1", "Feature2", "Feature3", "Feature4"],
                "Business_Model": "<Business_Model>",
                "Competition": "<Competition>",
                "Competitive_Advantage": "<Competitive_Advantage>"
            user
            {prompt_to_use}
            assistant
            """

        # Stream the output from the replicate model
        with st.spinner('Generating your unique app idea... This might take a moment 🚀'):
            response_text = ""
            event_texts = []
            for event in replicate.stream(
                "snowflake/snowflake-arctic-instruct",
                input={
                    "top_p": 0.9,
                    "prompt": prompt_to_use,
                    "temperature": 0.75,
                    "max_new_tokens": 1000,
                    "min_new_tokens": 0,
                    "prompt_template": prompt_template,
                    "presence_penalty": 1.15,
                    "frequency_penalty": 0.2
                },
            ):
                event_texts.append(str(event))
            
            # Join the event texts into a single response text
            response_text = ''.join(event_texts)

            try:
                # Parse the response text to extract JSON data
                json_data = json.loads(response_text)

                # Extract the primary color
                primary_color = json_data.get("Background_Color", "#FFFFFF")  # Default to white if not found

                # Apply custom CSS to change the background color of the main page
                st.markdown(
                    f"""
                    <style>
                    .stApp {{
                        background-color: {primary_color};
                    }}
                    </style>
                    """,
                    unsafe_allow_html=True
                )

                # Extracting details from the JSON data
                name = json_data.get("Name", "App Name")
                tagline = json_data.get("Tagline", "App Tagline")
                description = json_data.get("Description", "App Description")
                emoji = json_data.get("Emoji", "🎈")  # Default to a balloon emoji if not found
                problem = json_data.get("Problem", "")
                solution = json_data.get("Solution", "")
                features = json_data.get("Features", [])
                business_model = json_data.get("Business_Model", "")
                competition = json_data.get("Competition", "")
                competitive_advantage = json_data.get("Competitive_Advantage", "")

                # Check if the domain is available
                domain_name = name.split()[0].lower() + ".com"
                domain_available = not check_domain_exists(domain_name)
                domain_status = "The Domain is Available ✅" if domain_available else "Domain Taken ❌"

                # Format the markdown content
                markdown_content = f"""
                    # {name} {emoji}

                    **{tagline}**

                    {domain_status}

                    ### Description 📄
                    {description}

                    ### Problem 🛠️
                    {problem}

                    ### Solution 💡
                    {solution}

                    ### Features 🌟
                    {' 🌟 '.join(features)}

                    ### Business Model 💼
                    {business_model}

                    ### Competition 🏆
                    {competition}

                    ### Competitive Advantage 🎯
                    {competitive_advantage}
                    """


                # Display the markdown content
                st.markdown(markdown_content)

                # Make it rain emojis
                rain(
                    emoji=emoji,
                    font_size=36,
                    falling_speed=10,
                    animation_length=1,
                )

                # Function to create a DOCX file
                def create_docx(content, filename):
                    doc = Document()
                    doc.add_heading(name, 0)
                    doc.add_paragraph(tagline)
                    doc.add_paragraph(domain_status)
                    doc.add_heading('Description 📄', level=1)
                    doc.add_paragraph(description)
                    doc.add_heading('Problem 🛠️', level=1)
                    doc.add_paragraph(problem)
                    doc.add_heading('Solution 💡', level=1)
                    doc.add_paragraph(solution)
                    doc.add_heading('Features 🌟', level=1)
                    for feature in features:
                        doc.add_paragraph(f'• {feature}')
                    doc.add_heading('Business Model 💼', level=1)
                    doc.add_paragraph(business_model)
                    doc.add_heading('Competition 🏆', level=1)
                    doc.add_paragraph(competition)
                    doc.add_heading('Competitive Advantage 🎯', level=1)
                    doc.add_paragraph(competitive_advantage)

                    doc_output = BytesIO()
                    doc.save(doc_output)
                    doc_output.seek(0)

                    b64 = base64.b64encode(doc_output.read()).decode()
                    href = f'<a href="data:application/vnd.openxmlformats-officedocument.wordprocessingml.document;base64,{b64}" download="{filename}">Download {filename} 📄</a>'
                    return href

                # Button to download markdown as DOCX
                docx_link = create_docx(markdown_content, f"{name}.docx")
                st.markdown(docx_link, unsafe_allow_html=True)

            except json.JSONDecodeError:
                st.error("Failed to parse the JSON response. Please try again.")
            except Exception as e:
                st.error(f"An error occurred: {str(e)}")

# Apply the primary color if it is set in session state
if "primary_color" in st.session_state:
    primary_color = st.session_state.primary_color
    st.markdown(
        f"""
        <style>
        .stApp {{
            background-color: {primary_color};
        }}
        </style>
        """,
        unsafe_allow_html=True
    )
